import argparse
import json
import random
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

import re

def load_jsonl(path):
    items = []
    with open(path, "r", encoding="utf-8") as f:
        for i, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                items.append(obj)
            except json.JSONDecodeError:
                # Skip malformed lines silently; you can log if needed
                continue
    return items

def add_heading_styles(doc):
    # Make a slightly larger Normal for readability
    style = doc.styles["Normal"]
    font = style.font
    if font.size is None or font.size.pt < 11:
        font.size = Pt(11)

def write_item(doc, idx, item):
    paragraph = str(item.get("paragraph", "")).strip()
    paragraph = re.sub(r'-\s*\n\s*', '', paragraph)
    # paragraph = paragraph.replace("\n", " ").replace("\r", " ")
    
    question  = str(item.get("question", "")).strip()
    answer    = str(item.get("answer", "")).strip()

    # Title: Item N
    doc.add_heading(f"Item {idx}", level=2)

    # Paragraph
    p = doc.add_paragraph()
    p.add_run("Paragraph: ").bold = True
    doc.add_paragraph(paragraph) if paragraph else doc.add_paragraph("(missing)")

    # Question
    p = doc.add_paragraph()
    p.add_run("Question: ").bold = True
    doc.add_paragraph(question) if question else doc.add_paragraph("(missing)")

    # Answer
    p = doc.add_paragraph()
    p.add_run("Answer: ").bold = True
    doc.add_paragraph(answer) if answer else doc.add_paragraph("(missing)")

    # Spacer
    doc.add_paragraph()  # blank line

def main():
    ap = argparse.ArgumentParser(
        description="Sample N items from a .jsonl and export paragraph/question/answer to a Word .docx"
    )
    ap.add_argument("input_jsonl", type=Path, default="./04-Prototipar/4.1-Fuentes de Datos/4.1.1-GET_BAC_PUSH_BloB/downloads/dataQA.jsonl", help="Path to input .jsonl")
    ap.add_argument("-o", "--output", type=Path, default="./04-Prototipar/4.1-Fuentes de Datos/4.1.1-GET_BAC_PUSH_BloB/downloads/sampled_qas.docx",
                    help="Output .docx path (default: input filename with .docx)")
    ap.add_argument("-n", "--num", type=int, default=100, help="Number of items to sample")
    ap.add_argument("--seed", type=int, default=42, help="Random seed for reproducibility")
    args = ap.parse_args()

    items = load_jsonl(args.input_jsonl)
    if not items:
        raise SystemExit("No valid JSON objects found in the input file.")

    random.seed(args.seed)
    k = min(args.num, len(items))
    sampled = random.sample(items, k=k)

    out_path = args.output
    if out_path is None:
        out_path = args.input_jsonl.with_suffix(".docx")

    doc = Document()
    add_heading_styles(doc)

    # Document title
    doc.add_heading("Sampled Q&A Export", level=1)
    doc.add_paragraph(f"Source file: {args.input_jsonl.name}")
    doc.add_paragraph(f"Total items in source: {len(items)}")
    doc.add_paragraph(f"Sampled items: {k}")
    doc.add_paragraph()  # spacer

    for i, item in enumerate(sampled, 1):
        write_item(doc, i, item)
        # Optional: page break every 10 items for readability
        if i % 10 == 0 and i != k:
            # doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
            doc.add_page_break()

    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    main()


