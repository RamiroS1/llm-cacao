import requests
import tempfile
import pdfplumber
import fitz
from io import BytesIO

import dspy
import faiss
# If IVF/PQ:
gpu_index.nprobe = 8   # start 8→12→16 if recall needs it


import pickle

import pandas as pd
import re, math, time, html, unicodedata
from docx import Document

from pdf2image import convert_from_path
from difflib import SequenceMatcher

from paddleocr import PPStructureV3
from bs4 import BeautifulSoup
import pytesseract

import io
import numpy as np
import cv2
import os
import json

from sklearn.preprocessing import normalize

from sentence_transformers import SentenceTransformer
from sentence_transformers import util

from langchain.text_splitter import RecursiveCharacterTextSplitter
from signatures import UniversityRAG

import torch
torch.backends.cuda.matmul.allow_tf32 = True     # Ampere+
torch.backends.cudnn.allow_tf32 = True
torch.set_float32_matmul_precision("high")

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,        # total characters
    chunk_overlap=150,     # overlap between chunks
    separators=["\n\n", "\n", ".", " "]
)

def extract_tables(pdf, min_rows, min_cols):
    table_bboxes_per_page = []
    table_per_page = []

    for page in pdf.pages:
        valid_bboxes = []
        valid_tables = []
        tables = page.find_tables()

        for t in tables:
            extracted = t.extract()

            if not extracted:
                continue

            row_count = len(extracted)
            col_count = max(len(row) for row in extracted if row) if extracted else 0

            if row_count >= min_rows and col_count >= min_cols:
                valid_bboxes.append(t.bbox)
                valid_tables.append(extracted)

        # bboxes = [t.bbox for t in tables]
        table_bboxes_per_page.append(valid_bboxes)
        table_per_page.append(valid_tables)
    return table_bboxes_per_page, table_per_page

def table_to_text(table):
    """Flatten a table into a normalized text string."""
    return "\n".join(" | ".join(cell.strip().lower() for cell in row) for row in table)

def similar(a, b, threshold=0.9):
    """Return True if two strings are similar enough."""
    return SequenceMatcher(None, a, b).ratio() >= threshold

def rects_intersect(r1, r2):
    # r = [x0, y0, x1, y1]
    return not (r2[0] >= r1[2] or r2[2] <= r1[0] or r2[1] >= r1[3] or r2[3] <= r1[1])

def horizontal_overlap(b, x0, x1, p=0.5):
    overlap = max(0, min(b[2], x1) - max(b[0], x0))
    width = min(b[2] - b[0], x1 - x0)
    return overlap / width > p  # at least 50% overlap

def get_table_bboxes(url_path, min_rows=2, min_cols=2):

    with pdfplumber.open(url_path) as pdf:
        table_bboxes_per_page, table_per_page = extract_tables(pdf, min_rows, min_cols)

    return table_bboxes_per_page, table_per_page  # List of lists of [x0, top, x1, bottom]

def clean_table_markdown(raw_table):
    """
    Cleans and formats a markdown table string (with possible multi-line headers and rows).
    Returns a clean markdown-formatted table string.
    """
    lines = [line.strip() for line in raw_table.split('\n') if line.strip().startswith('|')]
    rows = [re.split(r'\s*\|\s*', line.strip('|')) for line in lines]

    # Detect header lines (stop at the separator row, e.g., |:----|)
    header_lines = []
    body_rows = []
    header_sep_found = False

    for i, row in enumerate(rows):
        if any(re.match(r'^:?-+:?$', cell.strip()) for cell in row):  # header separator
            header_sep_found = True
            break
        header_lines.append(row)
    body_start_idx = len(header_lines) + 1 if header_sep_found else 0
    body_rows = rows[body_start_idx:]

    if not header_lines:
        if rows:
            max_cols = max(len(r) for r in rows)
        else:
            max_cols = 4  # fallback default
        merged_header = [f"Column {i + 1}" for i in range(max_cols)]
    else:
        # Merge header columns vertically
        max_cols = max(len(r) for r in header_lines)
        padded_headers = [r + [''] * (max_cols - len(r)) for r in header_lines]
        merged_header = [' '.join(col).strip() for col in zip(*padded_headers)]

    """"# If no header was found, use default names
    if not merged_header or all(not h for h in merged_header):
        merged_header = [f"Column {i+1}" for i in range(max_cols)]"""

    # Merge multi-line rows in body
    merged_rows = []
    temp_row = [''] * max_cols

    for row in body_rows:
        row = row + [''] * (max_cols - len(row))  # pad row
        if any(cell.strip() for cell in row[1:]):  # New data row
            if any(temp_row[0].strip()):  # Append previous if not empty
                merged_rows.append(temp_row)
            temp_row = row
        else:
            temp_row[0] += ' ' + row[0]
    if any(temp_row[0].strip()):
        merged_rows.append(temp_row)

    # Build markdown table
    out_lines = []
    out_lines.append('| ' + ' | '.join(merged_header) + ' |')
    out_lines.append('|-' + '-|-'.join(['-' * len(h) for h in merged_header]) + '-|')
    for row in merged_rows:
        cells = [cell.strip() for cell in row]
        out_lines.append('| ' + ' | '.join(cells) + ' |')

    return '\n'.join(out_lines)

def extract_text_from_pdf(pdf_path, url=False, remove_h=True, header_height=50, footer_height=80):

    if url:
        try:
            res = requests.get(pdf_path, stream=True, timeout=10)
            res.raise_for_status()

            if 'application/pdf' not in res.headers.get('Content-Type', ''):
                print(f"Advertencia: {url} no es un PDF válido.")
                return ""

            table_bboxes, tables = get_table_bboxes(BytesIO(res.content))

        except requests.exceptions.RequestException as e:
            print(f"Error al descargar el PDF desde {url}: {e}")
            return "", ""
        except Exception as e:
            print(f"Error procesando el PDF desde {url}: {e}")
            return "", ""

        doc = fitz.open(stream=BytesIO(res.content), filetype="pdf")

    else:
        table_bboxes, tables = get_table_bboxes(pdf_path)
        doc = fitz.open(pdf_path)

    # df = pd.DataFrame(table[1:], columns=table[0])
    # markdown = df.to_markdown(index=False)

    full_tables = []
    for page_idx, tables_i in enumerate(tables):
        page = doc[page_idx]
        blocks = page.get_text("blocks")
        page_height = page.rect.height
        page_blocks_sorted = sorted(blocks, key=lambda b: b[1])

        for table_idx, table_i in enumerate(tables_i):
            df = pd.DataFrame(table_i[1:], columns=table_i[0])
            markdown = df.to_markdown(index=False)

            # Include a title for each table
            # Get table bounding box
            bbox = table_bboxes[page_idx][table_idx] if page_idx < len(table_bboxes) and table_idx < len(
                table_bboxes[page_idx]) else None

            # Find nearest block above the table as possible title
            title = ""
            if bbox:
                x0_table, y0_table, x1_table, y1_table = bbox
                candidates = [
                    b for b in page_blocks_sorted
                    if b[3] <= y0_table and # block is above the table
                       horizontal_overlap(b, x0_table, x1_table) #overlaps horizontally
                       # (x0_table < b[2] and x1_table > b[0])
                ]
                if candidates:
                    # Use the lowest (closest) block above the table
                    # best_block = sorted(candidates, key=lambda b: abs(b[3] - y0_table))[-1]
                    best_block = sorted(candidates, key=lambda b: abs(b[3] - y0_table))[0]
                    title = best_block[4].strip()

                    # Combine title and table

            markdown = clean_table_markdown(markdown)
            if title:
                full_tables.append(f"**{title}**\n\n{markdown}")
            else:
                full_tables.append(markdown)

    text = ""
    for i, page in enumerate(doc):

        if remove_h:
            page_text = ""
            page_height = page.rect.height
            blocks = page.get_text("blocks")

            bboxes_to_ignore = table_bboxes[i] if i < len(table_bboxes) else []

            for b in blocks:
                x0, y0, x1, y1, text_block, *_ = b

                if not (header_height < y0 < (page_height - footer_height)):
                    continue

                block_rect = [x0, y0, x1, y1]
                if any(rects_intersect(block_rect, bbox) for bbox in bboxes_to_ignore):
                    continue  # Skip table blocks

                page_text += text_block.strip() + "\n"

            '''filtered_text = [
                b[4] for b in blocks if header_height < b[1] < (page_height - footer_height)
            ]'''

            text += page_text + "\n"
        else:
            text += page.get_text("text")  # Extract text from each page
    return text, full_tables

def drop_fuzzy_duplicate_tables(tables, similarity_threshold=0.9):
    """
    Remove tables that are near-duplicates (fuzzy-matched) or contain boilerplate content.
    """
    kept = []
    seen_texts = []

    for table in tables:
        t_text = table_to_text(table)
        '''if is_probably_header_or_footer(t_text):
            continue'''

        # Check if similar to an already seen table
        if any(similar(t_text, seen, threshold=similarity_threshold) for seen in seen_texts):
            continue

        seen_texts.append(t_text)
        kept.append(table)

    return kept

def run_structure(table_engine, img):
    if hasattr(table_engine, "predict"):
        return table_engine.predict(img)   # V3
    if callable(table_engine):
        return table_engine(img)           # V2
    raise TypeError("PPStructure instance has neither .predict() nor __call__().")

def normalize_ppstruct_v3(result):
    """
    Convert PPStructureV3 result into a flat list of items like:
      {"type": "table", "bbox": [x0,y0,x1,y1], "res": {"html": "..."}}
      {"type": "text",  "bbox": [x0,y0,x1,y1], "text": "..."}

    Handles multiple possible field names safely.
    """
    items = []

    if not isinstance(result, dict):
        # Some configs may return a list directly (rare)
        result = {"items": result}

    # 1) Tables (preferred source)
    for t in result.get("tables", []) or []:
        if not isinstance(t, dict):
            continue
        bbox = t.get("bbox") or t.get("box") or t.get("poly")
        html = t.get("html") or (t.get("res") or {}).get("html")
        items.append({"type": "table", "bbox": bbox, "res": {"html": html}})

    # 2) Layout / OCR blocks (names vary)
    layout = result.get("layout") or result.get("ocr") or result.get("items") or []
    if isinstance(layout, dict):
        layout = layout.get("items", [])

    for it in layout or []:
        if not isinstance(it, dict):
            continue
        # normalize type/category
        itype = (it.get("type") or it.get("category") or it.get("label") or it.get("res_type") or "").lower()
        bbox = it.get("bbox") or it.get("box") or it.get("poly")
        # tables that appear inside layout
        if itype.startswith("table"):
            html = it.get("html") or (it.get("res") or {}).get("html")
            items.append({"type": "table", "bbox": bbox, "res": {"html": html}})
            continue
        # collect text-like items
        txt = it.get("text") or it.get("value") or it.get("content")
        if txt:
            items.append({"type": "text", "bbox": bbox, "text": txt})

    return items

def extract_text_from_scanned_pdf_2(pdf_path, remove_h=True, header_height=50, footer_height=80, from_url=False):
    text = ""
    full_tables = []
    _bbox_re = re.compile(r"bbox (\d+) (\d+) (\d+) (\d+)")
    max_gap = 400

    if from_url:
        response = requests.get(pdf_path)
        if response.status_code != 200:
            raise Exception(f"Failed to download PDF. Status code: {response.status_code}")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(response.content)
            pdf_path = tmp_file.name

    # Convert PDF pages to images
    # images = convert_from_path(pdf_path)
    # images = convert_from_path(pdf_path, dpi=400)
    images = convert_from_path(pdf_path, dpi=250, fmt="jpeg", size=(2200, None), thread_count=2, use_cropbox=True)

    # Initialize PPStructure for table extraction
    # table_engine = PPStructureV3(show_log=True, image_orientation=True)
    table_engine = PPStructureV3()

    lines_pre = []
    for img in images:
        # Convert PIL.Image to OpenCV (BGR) format

        # Extract text
        width, height = img.size

        if remove_h:
            img = img.crop((0, header_height, width, height - footer_height))

        tess_cfg = r'--oem 3 --psm 6 -l spa+eng'
        hocr = pytesseract.image_to_pdf_or_hocr(img, extension='hocr', config=tess_cfg)
        soup = BeautifulSoup(hocr, 'html.parser')

        nodes = soup.select("span.ocr_line, span.ocrx_line")
        lines = []
        for n in nodes:
            title = n.get("title", "")
            m = _bbox_re.search(title)
            if not m:
                continue
            x0, y0, x1, y1 = map(int, m.groups())
            txt = n.get_text(separator=" ", strip=True)
            if txt:
                lines.append((x0, y0, x1, y1, txt))

        blocks = []
        for block in soup.find_all('div', class_='ocr_carea'):
            title = block['title']
            coords = [int(v) for v in title.split(';')[0].replace('bbox ', '').split()]
            text_blocks = block.get_text(separator=' ', strip=True)
            blocks.append((coords[0], coords[1], coords[2], coords[3], text_blocks))

        '''# Run OCR with bounding boxes (Tesseract)
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        blocks = []
        for i in range(len(ocr_data['text'])):
            x, y, w, h = ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i]
            blocks.append((x, y, x + w, y + h, ocr_data['text'][i]))'''
        # text += pytesseract.image_to_string(img) + "\n"  # OCR text extraction

        # Convert PIL.Image to OpenCV
        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

        # result = table_engine(img_cv)
        result = run_structure(table_engine, img_cv)
        result = normalize_ppstruct_v3(result)

        bboxes_to_ignore = []
        for item in result:
            itype = item.get("type", "").lower()
            if itype == 'table':

                # Convert HTML → DataFrame
                try:
                    html_table = (item.get("res") or {}).get("html")

                    # Convert HTML → DataFrame
                    df = pd.read_html(io.StringIO(html_table), header=0)[0]
                    # df = pd.read_html(io.StringIO(html_table))[0]

                    # Output as Markdown
                    md_table = df.to_markdown(index=False)
                except Exception:
                    md_table = ""

                if md_table:
                    # Get table bbox
                    bbox = item.get("bbox") or [0, 0, 0, 0]
                    x0_table, y0_table, x1_table, y1_table = bbox

                    # Find nearest text above table
                    candidates = [
                        b for b in lines
                        if b[3] <= y0_table and horizontal_overlap(b, x0_table, x1_table, p=0.2)
                    ]
                    title = ""

                    if candidates:
                        # Closest block above table
                        candidates_sorted = sorted(candidates, key=lambda b: abs(b[3] - y0_table))

                        # Start from the closest line and grow upward while gaps are small
                        anchor = candidates_sorted[0]
                        x0a, y0a, x1a, y1a, _ = anchor
                        group = [anchor]

                        above = [b for b in candidates_sorted[1:] if b[3] <= y0a and abs(b[3] - group[-1][1]) <= max_gap]
                        block = sorted(group + above, key=lambda b: (b[1], b[0]))

                        #title = candidates_sorted[0][4].strip()
                        title = " ".join([b[4] for b in block]).strip()

                    # Table title from a later page
                    idx = 1
                    while len(title) < 250 and idx <= len(lines_pre):
                        title = lines_pre[-idx] + "\n" + title
                        idx += 1

                    bboxes_to_ignore.append((x0_table, y0_table, x1_table, y1_table))
                    md_table = clean_table_markdown(md_table)
                    if title:
                        full_tables.append(f"**{title}**\n\n{md_table}")
                    else:
                        full_tables.append(md_table)

        # text += page_text + "\n"

        # text
        # blocks_sorted = sorted(blocks, key=lambda b: (b[1], b[0]))
        lines_sorted = sorted(lines, key=lambda b: (b[1], b[0]))

        page_text = ""
        lines_pre = []
        for b in lines_sorted:
            x0, y0, x1, y1, text_line = b

            line_rect = [x0, y0, x1, y1]
            if any(rects_intersect(line_rect, bbox) for bbox in bboxes_to_ignore):
                continue  # Skip table blocks

            lines_pre.append(text_line)
            page_text += text_line.strip() + "\n"

        # Append full text from page
        # text += "\n".join([b[4] for b in blocks]) + "\n"
        text += page_text + "\n"
    return text, full_tables

def clean_text(text):
    # Remove page headers, footers, and TOC noise
    text = re.sub(r'Página\s+\d+\s+de\s+\d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Bucaramanga,?\s*\d{4}', '', text, flags=re.IGNORECASE)
    text = re.sub(r'Versión:\s*\d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'CONTENIDO.*?Pág\..*?\n', '', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'\.{5,}', '', text)  # remove dots used in TOC
    
    text = unicodedata.normalize("NFC", text)
    
    # Remove soft hyphen + common control chars
    text = text.replace("\u00AD", "")  # soft hyphen
    text = re.sub(r'[\u200B-\u200D\u2060]', '', text)  # zero-width
    
    # Join words broken by hyphen at line end: "fac-\n tor" -> "factor"
    text = re.sub(r'([A-Za-zÁÉÍÓÚÜÑáéíóúüñ])-\s*\n\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ])', r'\1\2', text)
    
    # Collapse noisy runs of punctuation/box-drawing/bullets/tilde lines, etc.
    text = re.sub(r'[.\,\-–—=~•··•\*•\^_\\\/\|\+…]{3,}', ' ', text)
    
    # Kill leftover “garbage clusters” that aren’t letters/digits/brackets/quotes/percent
    text = re.sub(r'[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9\(\)\[\]\{\}%\'"°ºª¿\?¡!\.,;:\s]', ' ', text)
    
    # Fix hyphenation that survived without newline: "crecimien- to" -> "crecimiento"
    text = re.sub(r'([A-Za-zÁÉÍÓÚÜÑáéíóúüñ])-\s+([a-záéíóúüñ])', r'\1\2', text)
    
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text.strip()

def clean_llm_output(text: str) -> str:
    # Remove leading/trailing square brackets, stray symbols, and extra whitespace
    text = text.strip()
    # text = re.sub(r'^[\[\]\-–•\d\.\s]+', '', text)  # remove leading symbols like [ ] -
    text = re.sub(r'^[\[\]\-–•\.\s]+', '', text)
    text = re.sub(r'[\[\]]+', '', text)            # remove all [ or ] inside
    text = re.sub(r'\s{2,}', ' ', text)            # reduce repeated spaces
    return text.strip()

def clean_output(text):
    # return re.sub(r"^\]?\s*", "", text)
    return re.sub(r"^[\]\s\n]+", "", text)

def split_by_paragraphs(text, min_length=100):
    text = clean_text(text)
    paragraphs = text_splitter.split_text(text)
    
    paragraphs = [p for p in paragraphs if len(p) >= min_length]
    return paragraphs

def is_content_chunk(chunk: str) -> bool:
    # Remove excess whitespace
    chunk_clean = chunk.strip()
    # If it’s too short or just numbers
    if len(chunk_clean) < 50:
        return False
    # If more than 50% is uppercase, likely a header or TOC
    upper_ratio = sum(1 for c in chunk_clean if c.isupper()) / len(chunk_clean)
    if upper_ratio > 0.5:
        return False
    # If it matches chapter/toc patterns
    '''if re.search(r'(TÍTULO|CAPÍTULO|[0-9]+\s*\.\s*[A-Z])', chunk_clean):
        return False'''
    # If it's mostly numbers or page refs
    if re.search(r'^\d+\s+[\w\s]*\d+$', chunk_clean):
        return False
    return True

def split_long_paragraphs(paragraphs, MAX_PARAGRAPH_LENGTH=2000):
    result = []
    for p in paragraphs:
        if len(p) > MAX_PARAGRAPH_LENGTH:
            # Split at sentence boundaries if needed
            parts = re.split(r'(?<=[.;:])\s+(?=[A-ZÁÉÍÓÚÑ])', p)
            temp = ''
            for part in parts:
                if len(temp) + len(part) < MAX_PARAGRAPH_LENGTH:
                    temp += ' ' + part
                else:
                    result.append(temp.strip())
                    temp = part
            if temp:
                result.append(temp.strip())
        else:
            result.append(p)
    return result

text_splitter1 = RecursiveCharacterTextSplitter(
    chunk_size=1200,        # total characters
    chunk_overlap=150,     # overlap between chunks
    separators=["\n\n", "\n", ".", " "]
)

def split_by_paragraphs1(text, min_length=100):
    #paragraphs = [p.strip() for p in text.split("\n\n") if len(p.strip()) > min_length]
    text = clean_text(text)
    paragraphs = text_splitter1.split_text(text)

    return paragraphs

"""def filter_relevant_paragraphs1(text, tables, query, model_match,  cutoff=0.8, from_parg=False):
    # paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 0 and len(p) > min_len]

    if from_parg:
        paragraphs = text
    else:
        paragraphs = split_by_paragraphs1(text)
        paragraphs = ["".join(pag.splitlines()) for pag in paragraphs]  # Remove "\n"

    # relevant = get_close_matches(query, paragraphs, n=3, cutoff=cutoff)  # cutoff=0.3

    paragraphs += tables"""

def filter_relevant_paragraphs1(text, tables, query, model_match,  cutoff=0.8, from_parg=False):
    # paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 0 and len(p) > min_len]

    if from_parg:
        paragraphs = text
    else:
        paragraphs = split_by_paragraphs1(text)
        paragraphs = ["".join(pag.splitlines()) for pag in paragraphs]  # Remove "\n"

    # relevant = get_close_matches(query, paragraphs, n=3, cutoff=cutoff)  # cutoff=0.3

    paragraphs += tables


    if paragraphs:
        query_embedding = model_match.encode(query, convert_to_tensor=True, normalize_embeddings=True)
        chunk_embeddings = model_match.encode(paragraphs, convert_to_tensor=True, normalize_embeddings=True)
        cosine_scores = util.cos_sim(query_embedding, chunk_embeddings)[0]
    else:
        cosine_scores = []

    # Filter by threshold
    relevant = []
    for idx, score in enumerate(cosine_scores):
        if score >= cutoff:
            # relevant.append((paragraphs[idx], float(score)))
            relevant.append(paragraphs[idx])

    return relevant


## Utils RAG
# Retriever component
class FaissRetriever:
    def __init__(self, index_path, doc_path):
        self.index = faiss.read_index(index_path)
        with open(doc_path, "rb") as f:
            self.documents = pickle.load(f)
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')


    def retrieve(self, query, k=3):
        query_embedding = self.embedder.encode([query])
        D, I = self.index.search(query_embedding, k)                                                                    
        return [self.documents[i] for i in I[0]]

def extract_score(text):
    # Use regex to find the first float-like number
    match = re.search(r"\b\d+\.\d+\b", text)
    return float(match.group()) if match else 0.0

class RelevanceScorer(dspy.Signature):
    """Predict the relevance of the document to answering the question with a score between 0 and 1."""
    question = dspy.InputField()
    document = dspy.InputField()
    relevance_score_to_answer_question = dspy.OutputField(desc="A number between 0 and 1 representing the relevance.")

# Reranked Retriever component
class RerankedFaissRetriever:
    def __init__(self, index_path, doc_path, reranker_model=None, k=10, model_match='intfloat/multilingual-e5-large'):
        self.index = faiss.read_index(index_path)
        with open(doc_path, "rb") as f:
            self.documents = pickle.load(f)
        # self.embedder = SentenceTransformer('all-MiniLM-L6-v2')
        # self.embedder = SentenceTransformer('all-mpnet-base-v2')
        # self.embedder = SentenceTransformer('intfloat/e5-large-v2')
        device = "cuda" if torch.cuda.is_available() else "cpu"
        self.embedder = SentenceTransformer(model_match, device=device)
        #self.embedder = self.embedder.half()
        
        self.k = k
        # self.reranker_model = reranker_model or dspy.Predict("question, document -> relevance_score")
        self.reranker_model = reranker_model or dspy.Predict(RelevanceScorer)
        
    
    def retrieve(self, query, top_n=3, len_context=5000, flag_embs=False):
        #model_match = SentenceTransformer('intfloat/multilingual-e5-large')
        
        cutoff_init = 0.8
        # Step 1: Use FAISS to get k
        query_embedding = normalize(self.embedder.encode([f"query: {query}"]))
        # query_embedding = normalize(self.embedder.encode([query]))
        D, I = self.index.search(query_embedding, self.k)
        candidates_ = [self.documents[i] for i in I[0]]
        paras = [pag["paragraph"] for pag in candidates_]

        # remove near-duplicates candidates
        # candidates_embedding = self.embedder.encode([pag['paragraph'] for pag in candidates_], convert_to_tensor=True)
        candidates_embedding = self.embedder.encode(
        [f"passage: {p}" for p in paras],
        convert_to_numpy=True,
        normalize_embeddings=True
    ).astype(np.float32)

        candidates = []
        seen_idx = set()
        for i, p in enumerate(candidates_):
            if any(util.cos_sim(candidates_embedding[i], candidates_embedding[j]) > 0.99 for j in seen_idx):
                continue
            seen_idx.add(i)
            candidates.append(p)

        scored = []
        for doc in candidates:
            response = self.reranker_model(question=query, document=f"passage: {doc['paragraph']}")
            score_text = response.relevance_score_to_answer_question
            score = extract_score(score_text) # Assuming score is numeric
            scored.append((doc, score))

        # Sort and return top_n
        ranked = sorted(scored, key=lambda x: x[1], reverse=True)
        """ranked_text = "\n\n".join(
            f"{doc[0]['division']}\n{doc[0]['paragraph']}" for doc in ranked[:top_n]
        )"""

        filtered_text = filter_relevant_paragraphs1(candidates, [], query, self.embedder , cutoff=cutoff_init,
                                                    from_parg=True)
        
        # print(filtered_text[0])
        
        ranked_text = "\n\n".join(
            f"{doc['subject']}\n{doc['paragraph']}" for doc in filtered_text
        )

        while len(ranked_text) > len_context:
            cutoff_init += 0.0005
            filtered_text = filter_relevant_paragraphs1(candidates, [], query, self.embedder , cutoff=cutoff_init,
                                                        from_parg=True)

            ranked_text = "\n\n".join(
                f"{doc['subject']}\n{doc['paragraph']}" for doc in filtered_text
            )
            
        if flag_embs:
            # Return both docs and their (normalized) embeddings for those docs
            # Convert embeddings to NumPy array of shape (len(filtered_text), d)
            matched_embs = np.array([
                candidates_embedding[paras.index(doc["paragraph"])]
                for doc in filtered_text
            ], dtype=np.float32)
            return filtered_text, matched_embs

        # return [doc for doc, _ in ranked[:top_n]]
        return filtered_text

def _cosine_sim(a, b):
    # a: (d,) or (1,d); b: (N,d)
    a = a / (np.linalg.norm(a, axis=-1, keepdims=True) + 1e-12)
    b = b / (np.linalg.norm(b, axis=-1, keepdims=True) + 1e-12)
    return (a @ b.T).reshape(-1)

class faster_UniversityRAGChain(dspy.Module):
    def __init__(
        self,
        retriever,
        model_match="intfloat/multilingual-e5-base",   # smaller than -large for speed
        top_n=8,               # fewer retrieved docs
        cutoff=0.75,           # similarity cutoff
        max_matched=6,         # cap how many matched docs we concatenate
        max_context_chars=5000 # hard cap for assembled context
    ):
        super().__init__()
        self.retriever = retriever
        self.model = dspy.Predict(UniversityRAG)
        
        device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"Loading sentence transformer for reranking on device: {device}")
        
        self.model_match = SentenceTransformer(model_match, device=device)
        self.model_match = self.model_match.half()
        
        if device == "cuda":
            try:
                self.model_match = self.model_match.half()
            except Exception:
                pass
            
        self.top_n = top_n
        self.cutoff = cutoff
        self.max_matched = max_matched
        self.max_context_chars = max_context_chars
        
        # Simple in-memory cache for doc embeddings by paragraph text
        self._emb_cache = {} 
    
    def _encode_query(self, q: str) -> np.ndarray:
        with torch.inference_mode():
            emb = self.model_match.encode(
                [f"query: {q}"],
                convert_to_numpy=True,
                batch_size=64,
                normalize_embeddings=True
            )[0]
        return emb.astype(np.float32)
    
    def _encode_paragraphs(self, paras):
    # Use cache when possible; batch-encode misses
        to_encode = [p for p in paras if p not in self._emb_cache]
        if to_encode:
            with torch.inference_mode():
                embs = self.model_match.encode(
                    [f"passage: {p}" for p in to_encode],
                    convert_to_numpy=True,
                    batch_size=128,
                    normalize_embeddings=True
                )
            for p, e in zip(to_encode, embs):
                self._emb_cache[p] = e.astype(np.float32)
        return np.stack([self._emb_cache[p] for p in paras], axis=0)
        
    def _assemble_context(self, docs):
    # join until max_context_chars is reached
        parts, total = [], 0
        for doc in docs:
            piece = f"{doc['subject']}\n{doc['paragraph']}"
            need = len(piece) + (2 if parts else 0)  # account for double newlines
            if total + need > self.max_context_chars:
                break
            parts.append(piece)
            total += need
        return "\n\n".join(parts) if parts else "No se encontró contexto relevante."
        
        
    def forward(self, question, ext_context):
        if ext_context:
            pre_answer = self.model(question=question, context=ext_context)
                
            ans_len = len(pre_answer.get("answer", ""))
            rea_len = len(pre_answer.get("reasoning", ""))
            allowance = max(1200, self.max_context_chars - (ans_len + min(rea_len, 1000)))
        else:
            pre_answer, allowance = None, self.max_context_chars
                
        # 2) Retrieve fewer docs quickly
        docs = self.retriever.retrieve(question, top_n=self.top_n, len_context=allowance)
        # docs, d_embs = self.retriever.retrieve(question, top_n=self.top_n, len_context=allowance, flag_embs=True)
    
        if not docs:
            context = "No se encontró contexto relevante."
            final = self.model(question=question, context=context)
            return final, context
            
        doc_texts = [doc['paragraph'].strip() for doc in docs]
            
        q_emb = self._encode_query(question.strip())
        d_embs = self._encode_paragraphs(doc_texts)
        
            
        #sims = _cosine_sim(q_emb[None, :], d_embs)   # (N,)
        sims = d_embs @ q_emb
        order = np.argsort(-sims)
        keep = []
        for idx in order:
            if sims[idx] < self.cutoff:
                continue
            keep.append((idx, sims[idx]))
            if len(keep) >= self.max_matched:
                break
                
        matched_docs = [docs[i] for i, _ in keep] if keep else [docs[order[0]]]
            
        # 4) Assemble compact context
        context = self._assemble_context(matched_docs)
            
        if pre_answer:
            addon = pre_answer.get("answer", "")
            rea = pre_answer.get("reasoning", "")
            extra = "Contexto a partir de Internet:\n" + (addon + ("\n" + rea if len(rea) <= 1000 else ""))
            # Fit the add-on within remaining budget
            rem = max(0, self.max_context_chars - len(context) - 2)
            if rem > 0 and extra:
                context = (context + "\n\n" + extra[:rem]) if context else extra[:rem]
            return self.model(question=question, context=context), context
            
        # 6) Normal path
        return self.model(question=question, context=context), context
    # key: str(paragraph), val: np.ndarray (d,)
        
        


    
class UniversityRAGChain(dspy.Module):
    def __init__(self, retriever, model_match):
        super().__init__()
        self.retriever = retriever
        self.model = dspy.Predict(UniversityRAG)
        self.model_match = SentenceTransformer(model_match)

    def forward(self, question, ext_context):

        if len(ext_context) > 0:
            pre_answwer = self.model(question=question, context=ext_context)

            if len(pre_answwer["reasoning"]) > 1000:
                len_context = 6000 - len(pre_answwer["answer"]) # default: 6000
            else:
                len_context = 6000 - len(pre_answwer["answer"]) - len(pre_answwer["reasoning"])
        else:
            len_context = 5000 # default: 5000

        docs = self.retriever.retrieve(question, top_n=10, len_context=len_context)

        if not docs:
            context = "No se encontró contexto relevante."
        else:
            # doc_texts = [doc['paragraph'] for doc in docs]
            doc_texts = [doc['paragraph'].lower() for doc in docs]

            # matched_paragraphs = get_close_matches(question, doc_texts, n=6, cutoff=0.15)  # 0.16

            matched_paragraphs = filter_relevant_paragraphs1(doc_texts, [], question.lower(), self.model_match, cutoff=0.8, from_parg=True)

            matched_docs = [doc for doc in docs if doc['paragraph'].lower() in matched_paragraphs]


            context = "\n\n".join(
                f"{doc['subject']}\n{doc['paragraph']}" for doc in matched_docs
            )
            # context = "\n\n".join

        if len(ext_context) > 0:
            if not docs:
                return pre_answwer, context
            else:
                if len(pre_answwer["reasoning"]) > 1000:
                    # context = "Internet: \n" + pre_answwer["answer"] + "\n\n" + context
                    context = context + "\n\n" + "Contexto a partir de Internet: \n" + pre_answwer["answer"]
                else:
                    # context = "Internet: \n" + pre_answwer["reasoning"] + "\n\n" + context
                    context = context + "\n\n" + "Contexto a partir de Internet: \n" + pre_answwer["answer"] + "\n" + pre_answwer["reasoning"]

        return self.model(question=question, context=context), context
    
    
def exists_qa(doc_stem, path_qa):
    if not os.path.exists(path_qa):
        return False
    with open(path_qa, "r", encoding="utf-8") as f:
        for line in f:
            try:
                entry = json.loads(line)
                if entry.get("doc_name") == doc_stem:
                    return True
            except json.JSONDecodeError:
                continue
    return False

def word_to_markdown(filepath):
    # Cargar archivo Word
    doc = Document(filepath)

    md_output = []  # Para las tablas
    txt_output = []  # Para el texto

    # Procesar párrafos y tablas en orden
    for element in doc.element.body:
        # Si es un párrafo
        if element.tag.endswith('p'):
            # Buscar el párrafo correspondiente
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:  # Solo agregar si no está vacío
                        # Agregar como texto simple al archivo .txt
                        txt_output.append(text)
                    break

        # Si es una tabla
        elif element.tag.endswith('tbl'):
            # Buscar la tabla correspondiente
            for table in doc.tables:
                if table._element == element:
                    # Convertir tabla a DataFrame
                    data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        data.append(row_data)

                    if data:
                        df = pd.DataFrame(data)

                        # Reemplazar vacíos
                        df = df.fillna("").astype(str).applymap(lambda x: x.strip())

                        # Eliminar filas completamente vacías
                        df = df[~(df.eq("").all(axis=1))]

                        # Eliminar columnas completamente vacías
                        df = df.loc[:, ~(df.eq("").all(axis=0))]

                        if not df.empty:
                            # Usar primera fila como encabezado si parece serlo
                            # o crear nombres ficticios
                            if len(df) > 1:
                                # Intentar usar primera fila como header
                                df.columns = df.iloc[0]
                                df = df[1:].reset_index(drop=True)
                            else:
                                df.columns = [f"col{i+1}" for i in range(df.shape[1])]

                            # Convertir a Markdown
                            md_table = df.to_markdown(index=False)
                            md_output.append(md_table)
                            # md_output.append("")  # Línea en blanco después de tabla
                    break

    """# Guardar archivo .md con las tablas
    with open(output_md, "w", encoding="utf-8") as f:
        f.write("\n\n".join(md_output))

    # Guardar archivo .txt con el texto
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write("\n\n".join(txt_output))

    print(f"✅ Tablas guardadas en: {output_md}")
    print(f"✅ Texto guardado en: {output_txt}")"""
    
    return "\n\n".join(txt_output), md_output
    
